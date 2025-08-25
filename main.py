import sys
import os
import shutil
import glob
import re
import json
import zipfile
import tarfile
from datetime import datetime

# Media/PDF/DOCX preview deps
import vlc
import fitz  # PyMuPDF
import docx  # python-docx

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QTreeWidget, QTreeWidgetItem, QTabWidget,
    QTextEdit, QToolBar, QAction, QFileDialog, QWidget, QHBoxLayout, QVBoxLayout,
    QTableWidgetItem, QStatusBar, QTabBar, QPushButton, QComboBox, QLabel,
    QScrollArea, QSplitter, QTableWidget, QLineEdit, QTableView, QRadioButton,
    QButtonGroup, QGroupBox, QMessageBox, QListWidget, QListWidgetItem, QFrame,
    QSlider
)
import piexif
from PIL import Image
from PyQt5.QtGui import QFont, QPixmap, QStandardItemModel, QStandardItem, QImage
from PyQt5.QtCore import Qt, QSize, QSortFilterProxyModel, QPropertyAnimation, QRect, QEasingCurve, QTimer

from ppadb.client import Client as AdbClient
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
from docx import Document

load_dotenv()

# ============================================================
# Shared regex & helpers
# ============================================================

USAGE_EVENT_RE = re.compile(r'time="([^"]+)"\s+type=([A-Z_]+)\s+package=([\w\.\d]+)(.*)')

def parse_usage_events(file_path):
    events = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                m = USAGE_EVENT_RE.search(line)
                if m:
                    events.append({
                        "time": m.group(1),
                        "event_type": m.group(2),
                        "package": m.group(3),
                        "extra_info": (m.group(4) or "").strip()
                    })
    except Exception:
        pass
    return events


# ============================================================
# PreviewWidget: unified preview area used in central tabs
# ============================================================

class PreviewWidget(QWidget):
    """
    A self-contained preview widget for one file path.
    Supports: image (jpg/png), text, pdf, docx, media (audio/video),
    and archives (zip, tar/tgz) with internal browser & sub-preview.
    """
    def __init__(self, path, temp_dir=None, parent=None):
        super().__init__(parent)
        self.path = path
        self.temp_dir = temp_dir or os.path.join(os.getcwd(), "TempData")
        os.makedirs(self.temp_dir, exist_ok=True)

        self.vlc_instance = None
        self.vlc_player = None
        self.slider = None
        self.timer = QTimer()
        self.timer.timeout.connect(self._update_slider)

        self.root = QVBoxLayout(self)
        self.root.setContentsMargins(8, 8, 8, 8)

        self._render(path)

    # ---------- lifecycle ----------

    def cleanup(self):
        try:
            if self.timer:
                self.timer.stop()
        except Exception:
            pass
        try:
            if self.vlc_player:
                self.vlc_player.stop()
                self.vlc_player.release()
                self.vlc_player = None
        except Exception:
            pass
        try:
            if self.vlc_instance:
                self.vlc_instance.release()
                self.vlc_instance = None
        except Exception:
            pass

    def closeEvent(self, event):
        self.cleanup()
        super().closeEvent(event)

    # ---------- rendering ----------

    def _render(self, path):
        ext = os.path.splitext(path)[1].lower()
        if ext in [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp"]:
            self._show_image(path)
        elif ext in [".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".mp4", ".avi", ".mkv", ".mov", ".wmv"]:
            self._show_media(path)
        elif ext in [".txt", ".log", ".csv", ".json", ".xml", ".html", ".md"]:
            self._show_text(path)
        elif ext == ".pdf":
            self._show_pdf(path)
        elif ext == ".docx":
            self._show_docx(path)
        elif ext in [".zip", ".tar", ".tgz", ".tar.gz"]:
            self._show_archive_browser(path)
        else:
            lbl = QLabel(f"Unsupported file type: {ext or '(none)'}")
            self.root.addWidget(lbl)

    # ---------- preview types ----------

    def _show_image(self, path):
        lbl = QLabel()
        lbl.setAlignment(Qt.AlignCenter)
        pixmap = QPixmap(path)
        lbl.setPixmap(pixmap.scaled(1000, 700, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        scroll = QScrollArea()
        scroll.setWidget(lbl)
        scroll.setWidgetResizable(True)
        self.root.addWidget(scroll)

    def _show_text(self, path):
        text = QTextEdit()
        text.setReadOnly(True)
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text.setText(f.read())
        except Exception as e:
            text.setText(f"Failed to open text: {e}")
        self.root.addWidget(text)

    def _show_pdf(self, path):
        scroll = QScrollArea()
        container = QWidget()
        vbox = QVBoxLayout(container)

        try:
            doc = fitz.open(path)
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                if pix.n < 4:  # RGB
                    img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format_RGB888)
                else:  # RGBA
                    img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format_RGBA8888)
                lbl = QLabel()
                lbl.setPixmap(QPixmap.fromImage(img))
                vbox.addWidget(lbl)
        except Exception as e:
            vbox.addWidget(QLabel(f"Failed to render PDF: {e}"))

        vbox.addStretch(1)
        scroll.setWidget(container)
        scroll.setWidgetResizable(True)
        self.root.addWidget(scroll)

    def _show_docx(self, path):
        text = QTextEdit()
        text.setReadOnly(True)
        try:
            d = docx.Document(path)
            content = "\n".join(p.text for p in d.paragraphs)
        except Exception as e:
            content = f"Failed to open DOCX: {e}"
        text.setText(content)
        self.root.addWidget(text)

    def _show_media(self, path):
        # Init VLC lazily
        try:
            self.vlc_instance = vlc.Instance()
            self.vlc_player = self.vlc_instance.media_player_new()
        except Exception as e:
            self.root.addWidget(QLabel(f"VLC init error: {e}"))
            return

        # Video surface
        video_frame = QFrame()
        video_frame.setFrameShape(QFrame.Box)
        video_frame.setMinimumSize(640, 360)
        self.root.addWidget(video_frame)

        # Bind VLC to QWidget window id
        wid = video_frame.winId()
        if sys.platform.startswith('linux'):
            self.vlc_player.set_xwindow(wid)
        elif sys.platform == "win32":
            self.vlc_player.set_hwnd(wid)
        elif sys.platform == "darwin":
            self.vlc_player.set_nsobject(int(wid))

        media = self.vlc_instance.media_new(path)
        self.vlc_player.set_media(media)

        # Controls
        ctrls = QWidget()
        h = QHBoxLayout(ctrls)
        play_btn = QPushButton("▶")
        pause_btn = QPushButton("⏸")
        stop_btn = QPushButton("⏹")

        play_btn.clicked.connect(lambda: self.vlc_player.play())
        pause_btn.clicked.connect(lambda: self.vlc_player.pause())
        def stop_and_reset():
            self.vlc_player.stop()
            if self.slider:
                self.slider.setValue(0)
        stop_btn.clicked.connect(stop_and_reset)

        self.slider = QSlider(Qt.Horizontal)
        self.slider.setRange(0, 1000)
        self.slider.sliderMoved.connect(lambda pos: self.vlc_player.set_position(pos/1000.0))

        h.addWidget(play_btn)
        h.addWidget(pause_btn)
        h.addWidget(stop_btn)
        h.addWidget(self.slider)
        self.root.addWidget(ctrls)

        # Start timer for slider
        self.timer.start(400)

    def _update_slider(self):
        try:
            if self.vlc_player and self.vlc_player.is_playing():
                pos = self.vlc_player.get_position()
                if pos > 0 and self.slider:
                    self.slider.setValue(int(pos * 1000))
        except Exception:
            pass

    # ---------- archive browser ----------

    def _show_archive_browser(self, path):
        """
        Left: list of members
        Right: sub-preview area using another PreviewWidget
        """
        splitter = QSplitter(Qt.Horizontal)
        lst = QListWidget()
        right_area = QWidget()
        right_layout = QVBoxLayout(right_area)
        right_layout.setContentsMargins(0, 0, 0, 0)

        splitter.addWidget(lst)
        splitter.addWidget(right_area)
        splitter.setStretchFactor(1, 1)
        self.root.addWidget(splitter)

        members = []
        ext = os.path.splitext(path)[1].lower()

        try:
            if ext == ".zip":
                with zipfile.ZipFile(path, "r") as z:
                    members = z.namelist()
            elif ext in [".tar", ".tgz", ".tar.gz"]:
                with tarfile.open(path, "r:*") as t:
                    members = t.getnames()
        except Exception as e:
            lst.addItem(f"(Failed to open archive: {e})")
            return

        for name in members:
            lst.addItem(name)

        def on_item(item):
            member = item.text()
            # extract to temp and preview
            base_dir = os.path.join(self.temp_dir, "ArchivePreview")
            os.makedirs(base_dir, exist_ok=True)
            # clear old files (optional: keep simple to avoid growth)
            # extract
            try:
                out_path = os.path.join(base_dir, os.path.basename(member))
                if ext == ".zip":
                    with zipfile.ZipFile(path, "r") as z:
                        z.extract(member, base_dir)
                else:
                    with tarfile.open(path, "r:*") as t:
                        t.extract(member, base_dir)

                # If the member was extracted into nested dirs, resolve the real file path:
                candidate = os.path.join(base_dir, member)
                if os.path.isdir(candidate):
                    # do nothing (cannot preview a folder)
                    self._replace_preview(right_layout, QLabel("Selected item is a directory"))
                    return
                else:
                    self._replace_preview(right_layout, PreviewWidget(candidate, temp_dir=self.temp_dir))
            except Exception as e:
                self._replace_preview(right_layout, QLabel(f"Extraction/preview failed: {e}"))

        lst.itemClicked.connect(on_item)

    def _replace_preview(self, layout, widget):
        while layout.count():
            item = layout.takeAt(0)
            w = item.widget()
            if w:
                try:
                    if isinstance(w, PreviewWidget):
                        w.cleanup()
                except Exception:
                    pass
                w.setParent(None)
                w.deleteLater()
        layout.addWidget(widget)


# ============================================================
# Usage Stats UI
# ============================================================

class ParameterFilterProxyModel(QSortFilterProxyModel):
    def __init__(self):
        super().__init__()
        self.filters = {}

    def set_filters(self, filters):
        self.filters = filters
        self.invalidateFilter()

    def filterAcceptsRow(self, source_row, source_parent):
        model = self.sourceModel()
        columns = {"time": 0, "event": 1, "package": 2, "extra": 3}
        for key, value in self.filters.items():
            if key in columns:
                idx = model.index(source_row, columns[key], source_parent)
                text = (model.data(idx) or "").lower()
                if value.lower() not in text:
                    return False
        return True


class UsageStatsWidget(QWidget):
    def __init__(self, adb_device, temp_root):
        super().__init__()
        self.device = adb_device
        self.out_dir = os.path.join(temp_root, "UsageStats")
        os.makedirs(self.out_dir, exist_ok=True)
        self.local_file = os.path.join(self.out_dir, "usage_dump.txt")

        self.model = QStandardItemModel()
        self.model.setHorizontalHeaderLabels(["Time", "Event Type", "Package", "Extra Info"])

        self.proxy = ParameterFilterProxyModel()
        self.proxy.setSourceModel(self.model)

        self.table = QTableView()
        self.table.setModel(self.proxy)
        self.table.setSortingEnabled(True)

        self.search_bar = QLineEdit()
        self.search_bar.setPlaceholderText("Search filters: time=2025-08 package=com.android.chrome event=RESUMED extra=...")
        self.search_bar.textChanged.connect(self.apply_filters)

        self.refresh_btn = QPushButton("Refresh")
        self.refresh_btn.clicked.connect(self.refresh_usage_stats)

        top = QHBoxLayout()
        top.addWidget(QLabel("Filter:"))
        top.addWidget(self.search_bar, 1)
        top.addWidget(self.refresh_btn)

        layout = QVBoxLayout(self)
        layout.addLayout(top)
        layout.addWidget(self.table)

        self.refresh_usage_stats()

    def refresh_usage_stats(self):
        try:
            self.device.shell('sh -c "dumpsys usagestats > /sdcard/usage_dump.txt"')
            self.device.pull("/sdcard/usage_dump.txt", self.local_file)
            events = parse_usage_events(self.local_file)
            self.populate_table(events)
        except Exception as e:
            self.model.removeRows(0, self.model.rowCount())
            self.model.appendRow([
                QStandardItem("ERROR"),
                QStandardItem("PULL_FAILED"),
                QStandardItem("adb/usagestats"),
                QStandardItem(str(e))
            ])

    def populate_table(self, events):
        self.model.removeRows(0, self.model.rowCount())
        for ev in events:
            self.model.appendRow([
                QStandardItem(ev["time"]),
                QStandardItem(ev["event_type"]),
                QStandardItem(ev["package"]),
                QStandardItem(ev["extra_info"])
            ])

    def apply_filters(self, text):
        filters = {}
        for part in text.split():
            if "=" in part:
                k, v = part.split("=", 1)
                filters[k.strip().lower()] = v.strip()
        self.proxy.set_filters(filters)


# ============================================================
# Settings, Sidebar, and Main Window
# ============================================================

class FixedWidthTabBar(QTabBar):
    def tabSizeHint(self, index):
        return QSize(220, self.sizeHint().height())


class ChatSidebar(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.main_app = parent
        self.setFixedWidth(0)
        self.setStyleSheet("""
            QWidget { background: #f6f7fb; }
            QTextEdit { background: white; border: 1px solid #ddd; }
            QLineEdit { background: white; border: 1px solid #ddd; padding: 4px; }
            QPushButton { padding: 6px 10px; }
        """)
        self.root = QVBoxLayout(self)
        self.root.setContentsMargins(10, 10, 10, 10)
        self.root.setSpacing(8)

        self._build_not_configured_card()
        self._build_chat_ui()
        self.show_not_configured()

    def _build_not_configured_card(self):
        self.not_configured_widget = QWidget()
        v = QVBoxLayout(self.not_configured_widget)
        v.setSpacing(8)
        v.addWidget(QLabel("<b>Agent not configured</b>"))
        msg = QLabel("The AI agent is not configured. Please configure it in Settings.")
        msg.setWordWrap(True)
        v.addWidget(msg)
        btn = QPushButton("Open Settings")
        btn.clicked.connect(self._open_settings_request)
        v.addWidget(btn)
        v.addStretch(1)

    def _build_chat_ui(self):
        self.chat_widget = QWidget()
        v = QVBoxLayout(self.chat_widget)
        v.setContentsMargins(0, 0, 0, 0)
        title = QLabel("Chat")
        title.setStyleSheet("font-weight: 600;")
        self.chat_history = QTextEdit()
        self.chat_history.setReadOnly(True)

        self.input_line = QLineEdit()
        self.input_line.setPlaceholderText("Type a message...")
        self.send_btn = QPushButton("Send")
        self.send_btn.clicked.connect(self.echo_message)

        row = QHBoxLayout()
        row.addWidget(self.input_line, 1)
        row.addWidget(self.send_btn)

        v.addWidget(title)
        v.addWidget(self.chat_history, 1)
        v.addLayout(row)

    def _open_settings_request(self):
        if self.main_app:
            self.main_app.show_settings_page()

    def echo_message(self):
        text = self.input_line.text().strip()
        if not text:
            return
        self.chat_history.append(f"<b>You:</b> {text}")
        self.chat_history.append(f"<b>AI:</b> Received: {text}")
        self.input_line.clear()

    def show_not_configured(self):
        self._set_content_widget(self.not_configured_widget)

    def show_chat_ui(self):
        self._set_content_widget(self.chat_widget)

    def _set_content_widget(self, widget):
        while self.root.count():
            item = self.root.takeAt(0)
            w = item.widget()
            if w:
                w.setParent(None)
        self.root.addWidget(widget)


class SettingsPage(QWidget):
    def __init__(self, main_app):
        super().__init__()
        self.main_app = main_app
        self.project_root = main_app.project_root
        self.config_path = main_app.config_path

        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(12, 12, 12, 12)

        theme_box = QGroupBox("Theme")
        tlay = QHBoxLayout()
        self.theme_dropdown = QComboBox()
        self._refresh_theme_list()
        self.apply_theme_btn = QPushButton("Apply Theme")
        self.apply_theme_btn.clicked.connect(self.apply_theme_clicked)
        tlay.addWidget(self.theme_dropdown, 1)
        tlay.addWidget(self.apply_theme_btn)
        theme_box.setLayout(tlay)
        layout.addWidget(theme_box)

        api_box = QGroupBox("AI Agent Configuration")
        apil = QVBoxLayout()

        mode_row = QHBoxLayout()
        self.online_radio = QRadioButton("Online (OpenAI API Key)")
        self.local_radio = QRadioButton("Local (LM Studio)")
        self.mode_group = QButtonGroup(self)
        self.mode_group.addButton(self.online_radio)
        self.mode_group.addButton(self.local_radio)
        mode_row.addWidget(self.online_radio)
        mode_row.addWidget(self.local_radio)
        apil.addLayout(mode_row)

        self.online_key_input = QLineEdit()
        self.online_key_input.setPlaceholderText("Enter OpenAI API key (sk-...)")
        apil.addWidget(self.online_key_input)

        local_group = QWidget()
        local_layout = QVBoxLayout(local_group)
        self.local_url_input = QLineEdit()
        self.local_url_input.setPlaceholderText("LM Studio base URL (e.g. http://127.0.0.1:8080)")
        self.local_model_input = QLineEdit()
        self.local_model_input.setPlaceholderText("Model name (e.g. gpt-4o or local model id)")
        local_layout.addWidget(self.local_url_input)
        local_layout.addWidget(self.local_model_input)
        apil.addWidget(local_group)

        self.status_label = QLabel("")
        apil.addWidget(self.status_label)

        api_box.setLayout(apil)
        layout.addWidget(api_box)

        btn_row = QHBoxLayout()
        self.save_btn = QPushButton("Save & Test Connection")
        self.save_btn.clicked.connect(self.save_and_test)
        self.back_btn = QPushButton("Back")
        self.back_btn.clicked.connect(self.on_back)
        btn_row.addStretch(1)
        btn_row.addWidget(self.save_btn)
        btn_row.addWidget(self.back_btn)
        layout.addLayout(btn_row)

        self.load_config_to_ui()
        self.online_radio.toggled.connect(self._update_field_visibility)
        self._update_field_visibility()

    def _refresh_theme_list(self):
        qss_dir = os.path.join(self.project_root, "QSS")
        files = []
        if os.path.isdir(qss_dir):
            for f in sorted(os.listdir(qss_dir)):
                if f.lower().endswith(".qss"):
                    files.append(f)
        self.theme_dropdown.clear()
        self.theme_dropdown.addItem("Default (None)")
        for f in files:
            self.theme_dropdown.addItem(f)

    def apply_theme_clicked(self):
        sel = self.theme_dropdown.currentText()
        if sel == "Default (None)":
            QApplication.instance().setStyleSheet("")
            self.main_app.current_theme = None
            self.main_app.statusBar.showMessage("Theme cleared")
            return
        qss_path = os.path.join(self.project_root, "QSS", sel)
        if os.path.exists(qss_path):
            try:
                with open(qss_path, "r", encoding="utf-8") as fh:
                    qss = fh.read()
                QApplication.instance().setStyleSheet(qss)
                self.main_app.current_theme = sel
                self.main_app.statusBar.showMessage(f"Applied theme: {sel}")
            except Exception as e:
                QMessageBox.warning(self, "Theme Error", f"Failed to apply theme: {e}")
        else:
            QMessageBox.warning(self, "Theme Error", "Selected theme file not found.")

    def _update_field_visibility(self):
        online = self.online_radio.isChecked()
        self.online_key_input.setVisible(online)
        self.local_url_input.setVisible(not online)
        self.local_model_input.setVisible(not online)

    def load_config_to_ui(self):
        cfg = {}
        try:
            if os.path.exists(self.config_path):
                with open(self.config_path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
        except Exception:
            cfg = {}
        mode = cfg.get("api_mode", "online")
        if mode == "local":
            self.local_radio.setChecked(True)
        else:
            self.online_radio.setChecked(True)

        self.online_key_input.setText(cfg.get("openai_key", ""))
        self.local_url_input.setText(cfg.get("host_url", ""))
        self.local_model_input.setText(cfg.get("model", ""))

        theme = cfg.get("theme")
        if theme:
            self._refresh_theme_list()
            idx = self.theme_dropdown.findText(theme)
            if idx >= 0:
                self.theme_dropdown.setCurrentIndex(idx)

        self._update_field_visibility()

    def save_and_test(self):
        api_mode = "local" if self.local_radio.isChecked() else "online"
        cfg = {
            "api_mode": api_mode,
            "openai_key": self.online_key_input.text().strip(),
            "host_url": self.local_url_input.text().strip(),
            "model": self.local_model_input.text().strip(),
            "theme": self.theme_dropdown.currentText() if self.theme_dropdown.currentText() != "Default (None)" else None
        }
        try:
            with open(self.config_path, "w", encoding="utf-8") as f:
                json.dump(cfg, f, indent=2)
        except Exception as e:
            QMessageBox.warning(self, "Save Error", f"Failed to save config: {e}")
            return

        if cfg.get("theme"):
            qss_path = os.path.join(self.project_root, "QSS", cfg["theme"])
            if os.path.exists(qss_path):
                try:
                    with open(qss_path, "r", encoding="utf-8") as fh:
                        QApplication.instance().setStyleSheet(fh.read())
                except Exception:
                    pass

        ok, message = self.main_app.test_connection(cfg)
        if ok:
            self.status_label.setText(f"<font color='green'>{message}</font>")
            self.main_app.statusBar.showMessage(message)
            self.main_app.on_config_updated(True, cfg)
        else:
            self.status_label.setText(f"<font color='red'>{message}</font>")
            self.main_app.statusBar.showMessage("API Connection Failed")
            self.main_app.on_config_updated(False, cfg)

    def on_back(self):
        self.main_app.hide_settings_page()


# ============================================================
# Main Window
# ============================================================

class DroidForen(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DroidForen")
        self.setGeometry(200, 100, 1200, 720)
        self.setFont(QFont("Segoe UI", 10))

        self.device = None
        self.devices_map = {}
        self._chat_open = False

        # WhatsApp removed completely from file-types
        self.ext_map = {
            "Photos": {".jpg", ".jpeg", ".png"},
            "Documents": {".pdf", ".doc", ".docx", ".txt", ".xls", ".xlsx", ".ppt", ".pptx"},
            "Videos": {".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv"},
            "Audio": {".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a"},
            "Archives": {".zip", ".rar", ".7z", ".tar", ".gz", ".tgz", ".bz2", ".xz", ".iso"}
        }

        self.SectionList = [
            "Call Logs", "SMS", "Photos", "Videos",
            "Audio", "Documents", "Contacts", "Archives", "Usage Stats"
        ]

        self.project_root = os.path.dirname(os.path.abspath(__file__))
        self.temp_dir = os.path.join(self.project_root, "TempData")
        os.makedirs(self.temp_dir, exist_ok=True)

        self.config_path = os.path.join(self.project_root, "config.json")
        self.current_theme = None

        self.central_container = QWidget()
        self.setCentralWidget(self.central_container)
        self.main_layout = QVBoxLayout(self.central_container)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.content_layout = QHBoxLayout()
        self.main_layout.addLayout(self.content_layout)

        self.device_dropdown = QComboBox()
        self.device_dropdown.setPlaceholderText("Select a device")
        self.main_layout.addWidget(self.device_dropdown, alignment=Qt.AlignCenter)
        self.device_dropdown.setVisible(False)

        self.connect_button = QPushButton("Connect")
        self.connect_button.clicked.connect(self.connect_device)
        self.main_layout.addWidget(self.connect_button, alignment=Qt.AlignCenter)

        self.sidebarTree = QTreeWidget()
        self.sidebarTree.setHeaderHidden(True)
        self.sidebarTree.setMaximumWidth(260)
        self.sidebarTree.itemClicked.connect(self.open_or_focus_tab)
        self.content_layout.addWidget(self.sidebarTree)

        self.previewTabs = QTabWidget()
        self.previewTabs.setTabsClosable(True)
        self.previewTabs.tabCloseRequested.connect(self._close_tab_cleanup)
        self.previewTabs.setMovable(True)
        self.previewTabs.setTabBar(FixedWidthTabBar())
        self.content_layout.addWidget(self.previewTabs, 1)

        self.settings_page = SettingsPage(self)
        self.settings_page.setVisible(False)
        self.content_layout.addWidget(self.settings_page, 1)

        self.toolbar = QToolBar("Main Toolbar")
        self.toolbar.setMovable(False)
        self.addToolBar(Qt.TopToolBarArea, self.toolbar)

        open_act = QAction("Open File", self)
        open_act.triggered.connect(self.open_file_dialog)
        export_act = QAction("Export", self)
        export_act.triggered.connect(self.export_data)
        settings_act = QAction("Settings", self)
        settings_act.triggered.connect(self.show_settings_page)
        generate_report_act = QAction("Generate Report", self)
        generate_report_act.triggered.connect(self.generate_report)
        disconnect_act = QAction("Disconnect", self)
        disconnect_act.triggered.connect(self.disconnect_device)

        self.toolbar.addAction(open_act)
        self.toolbar.addAction(export_act)
        self.toolbar.addAction(settings_act)
        self.toolbar.addAction(generate_report_act)
        self.toolbar.addAction(disconnect_act)

        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.statusBar.showMessage("Not Connected")

        self.sidebarTree.setVisible(False)
        self.previewTabs.setVisible(False)
        self.toolbar.setVisible(False)
        self.statusBar.setVisible(False)

        self.populate_list()

        # Chat sidebar
        self.chatSidebar = ChatSidebar(self)
        self.chat_width = 320
        self.chatSidebar.setFixedHeight(self.height())
        self.chatSidebar.move(self.width(), 0)

        self.chatToggleBtn = QPushButton("<", self)
        self.chatToggleBtn.setFixedWidth(28)
        self.chatToggleBtn.setFixedHeight(80)
        self.chatToggleBtn.clicked.connect(self.toggle_chat_sidebar)
        self.chatToggleBtn.raise_()
        self.position_chat_controls()

        self.chatAnim = QPropertyAnimation(self.chatSidebar, b"geometry")
        self.chatAnim.setDuration(250)
        self.chatAnim.setEasingCurve(QEasingCurve.InOutCubic)

        self.loaded_config = {}
        self.load_config_and_test_on_startup()

    # ---------------------------- Chat layout helpers ----------------------------

    def position_chat_controls(self):
        self.chatToggleBtn.move(self.width() - self.chatToggleBtn.width(), int((self.height() - self.chatToggleBtn.height()) / 2))
        if not self._chat_open:
            self.chatSidebar.move(self.width(), 0)
        else:
            self.chatSidebar.move(self.width() - self.chat_width, 0)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.chatSidebar.setFixedHeight(self.height())
        self.position_chat_controls()

    def toggle_chat_sidebar(self):
        self.chatSidebar.setFixedWidth(self.chat_width)
        self.chatAnim.stop()
        if self._chat_open:
            start = QRect(self.width() - self.chat_width, 0, self.chat_width, self.height())
            end = QRect(self.width(), 0, self.chat_width, self.height())
            self.chatAnim.setStartValue(start)
            self.chatAnim.setEndValue(end)
            self.chatAnim.start()
            self._chat_open = False
            self.chatToggleBtn.setText("<")
        else:
            start = QRect(self.width(), 0, self.chat_width, self.height())
            end = QRect(self.width() - self.chat_width, 0, self.chat_width, self.height())
            self.chatAnim.setStartValue(start)
            self.chatAnim.setEndValue(end)
            self.chatAnim.start()
            self._chat_open = True
            self.chatToggleBtn.setText(">")

    # ---------------------------- Device & Evidence ----------------------------

    def populate_list(self):
        try:
            client = AdbClient(host="127.0.0.1", port=5037)
            devices = client.devices()
            self.devices_map = {d.serial: d for d in devices}
            self.device_dropdown.clear()

            if not devices:
                self.statusBar.setVisible(True)
                self.statusBar.showMessage("No devices found.")
                self.device_dropdown.setVisible(False)
                return

            self.device_dropdown.addItems([d.serial for d in devices])
            self.device_dropdown.setVisible(True)
            self.statusBar.setVisible(True)
            self.statusBar.showMessage("Select a device and click Connect")

        except Exception as e:
            self.statusBar.setVisible(True)
            self.statusBar.showMessage(f"Error fetching devices: {e}")

    def connect_device(self):
        deviceselect = self.device_dropdown.currentText().strip()
        if not deviceselect:
            self.statusBar.showMessage("Please select a device first.")
            return
        try:
            client = AdbClient(host="127.0.0.1", port=5037)
            live_devices = {d.serial: d for d in client.devices()}

            if deviceselect not in live_devices:
                self.statusBar.showMessage("Selected device is no longer connected.")
                self.populate_list()
                return

            self.device = live_devices[deviceselect]
            self.statusBar.showMessage(f"Connected to {self.device.serial}")

            info = {
                "Model": self.device.shell("getprop ro.product.model").strip(),
                "Manufacturer": self.device.shell("getprop ro.product.manufacturer").strip(),
                "Android Version": self.device.shell("getprop ro.build.version.release").strip(),
                "Device Name": self.device.shell("getprop ro.product.device").strip(),
                "Serial Number": self.device.shell("getprop ro.serialno").strip(),
                "CPU ABI": self.device.shell("getprop ro.product.cpu.abi").strip()
            }

            self.sidebarTree.clear()
            device_root = QTreeWidgetItem(["Connected Device"])
            for key, value in info.items():
                child = QTreeWidgetItem([f"{key}: {value}"])
                device_root.addChild(child)
            self.sidebarTree.addTopLevelItem(device_root)
            self.sidebarTree.addTopLevelItems(QTreeWidgetItem([section]) for section in self.SectionList)
            self.sidebarTree.setVisible(True)
            self.previewTabs.setVisible(True)
            self.toolbar.setVisible(True)
            self.statusBar.setVisible(True)
            self.connect_button.setVisible(False)
            self.device_dropdown.setVisible(False)

            try:
                self.collect_evidence(info)
            except Exception as e:
                self.statusBar.showMessage(f"Connected but evidence collection failed: {e}")

        except Exception as e:
            self.statusBar.setVisible(True)
            self.statusBar.showMessage(f"Connection failed: {e}")

    def disconnect_device(self):
        # Cleanup tabs (VLC etc.)
        self._close_all_tabs_cleanup()
        shutil.rmtree(self.temp_dir, ignore_errors=True)
        os.makedirs(self.temp_dir, exist_ok=True)
        self.sidebarTree.clear()
        self.previewTabs.clear()
        self.sidebarTree.setVisible(False)
        self.previewTabs.setVisible(False)
        self.toolbar.setVisible(False)
        self.statusBar.setVisible(False)
        self.connect_button.setVisible(True)
        self.device_dropdown.setVisible(True)
        self.statusBar.showMessage("Disconnected")
        self.populate_list()
        self.cleanup_evidence()

    def evidence_file_path(self):
        return os.path.join(self.project_root, "evidence.json")

    def collect_evidence(self, device_info=None):
        self.statusBar.showMessage("Collecting evidence...")
        evidence = {}
        try:
            if device_info is None:
                device_info = {
                    "Model": self.device.shell("getprop ro.product.model").strip(),
                    "Manufacturer": self.device.shell("getprop ro.product.manufacturer").strip(),
                    "Android Version": self.device.shell("getprop ro.build.version.release").strip(),
                    "Device Name": self.device.shell("getprop ro.product.device").strip(),
                    "Serial Number": self.device.shell("getprop ro.serialno").strip(),
                    "CPU ABI": self.device.shell("getprop ro.product.cpu.abi").strip()
                }
        except Exception as e:
            device_info = {"error": f"Failed to read device props: {e}"}

        calls = self._collect_call_logs()
        sms = self._collect_sms()
        files = self._collect_files_summary(limit=200)
        usage_stats = self._collect_usage_stats()

        case_id = f"CASE-{device_info.get('Serial Number','unknown')}-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        evidence["device_info"] = {"case_id": case_id, **device_info}
        evidence["files"] = files
        evidence["calls"] = calls
        evidence["sms"] = sms
        evidence["usage_stats"] = usage_stats

        try:
            with open(self.evidence_file_path(), "w", encoding="utf-8") as f:
                json.dump(evidence, f, indent=2)
            self.statusBar.showMessage(f"Evidence saved to {self.evidence_file_path()}")
        except Exception as e:
            self.statusBar.showMessage(f"Failed to save evidence.json: {e}")

    def cleanup_evidence(self):
        try:
            ef = self.evidence_file_path()
            if os.path.exists(ef):
                os.remove(ef)
                self.statusBar.showMessage("evidence.json deleted")
        except Exception as e:
            self.statusBar.showMessage(f"Failed to remove evidence.json: {e}")

    def _collect_call_logs(self):
        try:
            raw = self.device.shell("content query --uri content://call_log/calls")
            entries = []
            for block in raw.split("Row"):
                if not block.strip():
                    continue
                entry = {}
                for part in block.split(","):
                    if "=" in part:
                        k, v = part.strip().split("=", 1)
                        entry[k.strip()] = v.strip()
                if entry:
                    entries.append(entry)
            return entries
        except Exception as e:
            return [{"error": str(e)}]

    def _collect_sms(self):
        try:
            raw = self.device.shell("content query --uri content://sms/")
            entries = []
            for block in raw.split("Row"):
                if not block.strip():
                    continue
                entry = {}
                for part in block.split(","):
                    if "=" in part:
                        k, v = part.strip().split("=", 1)
                        entry[k.strip()] = v.strip()
                if entry:
                    entries.append(entry)
            return entries
        except Exception as e:
            return [{"error": str(e)}]

    def _collect_files_summary(self, limit=200):
        try:
            paths_to_scan = ["/sdcard/DCIM", "/sdcard/Pictures", "/sdcard/Documents", "/sdcard/Download", "/sdcard/Movies", "/sdcard/Music"]
            files = []
            count = 0
            for base in paths_to_scan:
                try:
                    raw = self.device.shell(f"ls -R {base}")
                except Exception:
                    continue
                current_dir = base
                for line in raw.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    if line.endswith(":"):
                        current_dir = line[:-1].strip()
                        if not current_dir.startswith("/"):
                            current_dir = f"/{current_dir}"
                        continue
                    for part in line.split():
                        lower = part.lower()
                        for cat, exts in self.ext_map.items():
                            if any(lower.endswith(ext) for ext in exts):
                                files.append({"path": f"{current_dir}/{part}", "type": cat})
                                count += 1
                                break
                        if count >= limit:
                            break
                    if count >= limit:
                        break
                if count >= limit:
                    break
            return files
        except Exception as e:
            return [{"error": str(e)}]

    def _collect_usage_stats(self):
        try:
            out_dir = os.path.join(self.temp_dir, "UsageStats")
            os.makedirs(out_dir, exist_ok=True)
            self.device.shell('sh -c "dumpsys usagestats > /sdcard/usage_dump.txt"')
            local_file = os.path.join(out_dir, "usage_dump.txt")
            self.device.pull("/sdcard/usage_dump.txt", local_file)
            return parse_usage_events(local_file)
        except Exception as e:
            return [{"error": f"usage_stats: {e}"}]

    # ---------------------------- Central Preview Helpers ----------------------------

    def _open_preview_tab(self, title, path):
        """
        Opens given 'path' in a new tab using the unified PreviewWidget.
        """
        # Focus existing tab if same title is open
        for i in range(self.previewTabs.count()):
            if self.previewTabs.tabText(i) == title:
                self.previewTabs.setCurrentIndex(i)
                return

        widget = PreviewWidget(path, temp_dir=self.temp_dir)
        idx = self.previewTabs.addTab(widget, title)
        self.previewTabs.setCurrentIndex(idx)

    def _close_tab_cleanup(self, index):
        w = self.previewTabs.widget(index)
        try:
            if isinstance(w, PreviewWidget):
                w.cleanup()
        except Exception:
            pass
        self.previewTabs.removeTab(index)

    def _close_all_tabs_cleanup(self):
        for i in reversed(range(self.previewTabs.count())):
            self._close_tab_cleanup(i)

    # ---------------------------- UI actions ----------------------------

    def open_or_focus_tab(self, item):
        title = item.text(0)
        parent = item.parent()

        # File click under media/doc trees -> preview
        file_sections = {"Photos", "Videos", "Audio", "Documents", "Archives"}
        if parent and parent.text(0) in file_sections:
            # The file was pulled earlier into TempData/<Section>/
            local_path = os.path.join(self.temp_dir, parent.text(0), title)
            if os.path.exists(local_path):
                # Use unified preview
                self._open_preview_tab(title, local_path)
            else:
                self.open_tab("Error", f"File not found locally: {local_path}")
            return

        # Top-level tabs
        for i in range(self.previewTabs.count()):
            if self.previewTabs.tabText(i) == title:
                self.previewTabs.setCurrentIndex(i)
                return

        if title == "Call Logs":
            self.show_call_logs()
        elif title == "SMS":
            content = self.device.shell("content query --uri content://sms/")
            self.open_tab(title, content)
        elif title == "Contacts":
            content = self.device.shell("content query --uri content://contacts/phones/")
            self.open_tab(title, content)
        elif title == "Usage Stats":
            usage_widget = UsageStatsWidget(self.device, self.temp_dir)
            idx = self.previewTabs.addTab(usage_widget, "Usage Stats")
            self.previewTabs.setCurrentIndex(idx)
        elif title in file_sections:
            # Populate (or refresh) list of files under this section
            self.Extract(title)
        else:
            # Any other text node (like "Connected Device" children)
            pass

    def show_call_logs(self):
        try:
            raw = self.device.shell("content query --uri content://call_log/calls")
            entries = [e.strip() for e in raw.split("Row") if e.strip()]
            headers = ["Name", "Number", "Type", "Date", "Duration"]
            table = QTableWidget()
            table.setColumnCount(len(headers))
            table.setHorizontalHeaderLabels(headers)
            table.setRowCount(len(entries))
            for row_idx, entry in enumerate(entries):
                entry_dict = {}
                for part in entry.split(","):
                    if "=" in part:
                        key, val = part.strip().split("=", 1)
                        val = val.strip()
                        if val in ("NULL", ""):
                            val = "N/A"
                        entry_dict[key.strip()] = val

                name = entry_dict.get("name", "N/A")
                number = entry_dict.get("number", "N/A")
                call_type = self.call_type(entry_dict.get("type", "0"))
                date = self.format_date(entry_dict.get("date", "0"))
                duration = f"{entry_dict.get('duration', '0')} sec"

                table.setItem(row_idx, 0, QTableWidgetItem(name))
                table.setItem(row_idx, 1, QTableWidgetItem(number))
                table.setItem(row_idx, 2, QTableWidgetItem(call_type))
                table.setItem(row_idx, 3, QTableWidgetItem(date))
                table.setItem(row_idx, 4, QTableWidgetItem(duration))

            table.resizeColumnsToContents()
            index = self.previewTabs.addTab(table, "Call Logs")
            self.previewTabs.setCurrentIndex(index)

        except Exception as e:
            self.open_tab("Call Logs", f"Failed to load call logs: {e}")

    def open_tab(self, title, content):
        for i in range(self.previewTabs.count()):
            if self.previewTabs.tabText(i) == title:
                self.previewTabs.setCurrentIndex(i)
                return
        editor = QTextEdit()
        editor.setText(content)
        editor.setReadOnly(True)
        index = self.previewTabs.addTab(editor, title)
        self.previewTabs.setCurrentIndex(index)

    def call_type(self, call_type):
        mapping = {
            "1": "Incoming",
            "2": "Outgoing",
            "3": "Missed",
            "4": "Voicemail",
            "5": "Rejected",
            "6": "Blocked",
            "7": "Answered Externally"
        }
        return mapping.get(call_type, "Unknown")

    def format_date(self, timestamp):
        try:
            ts = int(timestamp)
            return datetime.fromtimestamp(ts / 1000).strftime("%Y-%m-%d %H:%M:%S")
        except:
            return "Invalid"

    def open_file_dialog(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "All Files (*)")
        if not file_path:
            return
        # Open external file in central preview tab
        title = os.path.basename(file_path)
        self._open_preview_tab(title, file_path)

    # ---------------------------- Extraction & Export ----------------------------

    def Extract(self, section):
        """
        Scans typical device dirs and pulls matching files to temp/section,
        then populates the sidebar tree children under the section.
        """
        try:
            temp_sub_dir = os.path.join(self.temp_dir, section)
            os.makedirs(temp_sub_dir, exist_ok=True)
            # Cleanup existing children files in temp for that section
            for f in glob.glob(os.path.join(temp_sub_dir, "*")):
                try:
                    os.remove(f)
                except:
                    pass

            # Scan /sdcard recursively (like previous logic)
            raw = self.device.shell("ls -R /sdcard")
            lines = raw.splitlines()
            current_dir = "/sdcard"
            file_paths = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.endswith(":"):
                    current_dir = line[:-1].strip()
                    if not current_dir.startswith("/"):
                        current_dir = f"/{current_dir}"
                    continue
                for part in line.split():
                    if any(part.lower().endswith(ext) for ext in self.ext_map.get(section, set())):
                        file_paths.append(f"{current_dir}/{part}")

            downloaded = []
            for path in file_paths:
                filename = os.path.basename(path)
                dest = os.path.join(temp_sub_dir, filename)
                try:
                    self.device.pull(path, dest)
                    downloaded.append(dest)
                except:
                    pass

            # Populate the section nodes with local file names
            for i in range(self.sidebarTree.topLevelItemCount()):
                item = self.sidebarTree.topLevelItem(i)
                if item.text(0) == section:
                    item.takeChildren()
                    for file_path in downloaded:
                        child = QTreeWidgetItem([os.path.basename(file_path)])
                        item.addChild(child)
                    item.setExpanded(True)
                    break

        except Exception as e:
            self.open_tab(section, f"Error loading {section}: {e}")

    def export_data(self):
        current_tab_title = None
        if self.previewTabs.count() > 0 and self.previewTabs.currentIndex() >= 0:
            current_tab_title = self.previewTabs.tabText(self.previewTabs.currentIndex())
        folder = QFileDialog.getExistingDirectory(self, "Select export folder")
        if not folder:
            return
        try:
            if current_tab_title == "Photos":
                self.device.shell("ls -R /sdcard/Pictures/")
                self.device.pull("/sdcard/Pictures/", folder)
            elif current_tab_title == "Call Logs":
                data = self.device.shell("content query --uri content://call_log/calls")
                with open(os.path.join(folder, "call_logs.txt"), "w", encoding="utf-8") as f:
                    f.write(data)
            elif current_tab_title == "SMS":
                data = self.device.shell("content query --uri content://sms/")
                with open(os.path.join(folder, "sms.txt"), "w", encoding="utf-8") as f:
                    f.write(data)
            elif current_tab_title == "Contacts":
                data = self.device.shell("content query --uri content://contacts/phones/")
                with open(os.path.join(folder, "contacts.txt"), "w", encoding="utf-8") as f:
                    f.write(data)
            elif current_tab_title == "Usage Stats":
                src = os.path.join(self.temp_dir, "UsageStats", "usage_dump.txt")
                if os.path.exists(src):
                    shutil.copy2(src, os.path.join(folder, "usage_dump.txt"))
                else:
                    self.statusBar.showMessage("No local usage_dump.txt to export.")
                    return
            else:
                self.statusBar.showMessage("No export action for this tab.")
                return

            self.statusBar.showMessage(f"Exported {current_tab_title} to {folder}")
        except Exception as e:
            self.statusBar.showMessage(f"Export failed: {str(e)}")

    # ---------------------------- Settings / LLM ----------------------------

    def show_settings_page(self):
        self.previewTabs.setVisible(False)
        self.settings_page.setVisible(True)
        self.sidebarTree.setVisible(True)

    def hide_settings_page(self):
        self.settings_page.setVisible(False)
        self.previewTabs.setVisible(True)

    def load_config_and_test_on_startup(self):
        cfg = {}
        try:
            if os.path.exists(self.config_path):
                with open(self.config_path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
        except Exception:
            cfg = {}

        self.loaded_config = cfg
        if cfg:
            ok, message = self.test_connection(cfg)
            if ok:
                self.statusBar.setVisible(True)
                self.statusBar.showMessage(message)
                self.chatSidebar.show_chat_ui()
                theme = cfg.get("theme")
                if theme:
                    qss_path = os.path.join(self.project_root, "QSS", theme)
                    if os.path.exists(qss_path):
                        try:
                            with open(qss_path, "r", encoding="utf-8") as fh:
                                QApplication.instance().setStyleSheet(fh.read())
                        except Exception:
                            pass
                self.current_theme = theme
                return
        self.chatSidebar.show_not_configured()
        self.statusBar.setVisible(True)
        if not cfg:
            self.statusBar.showMessage("Agent not configured")
        else:
            self.statusBar.showMessage("Agent connection failed")

    def test_connection(self, cfg):
        if ChatOpenAI is None:
            return False, "ChatOpenAI import not available (install langchain-openai)"
        try:
            mode = cfg.get("api_mode", "online")
            model = cfg.get("model") or "gpt-4o"
            if mode == "local":
                url = cfg.get("host_url", "").strip()
                if not url:
                    return False, "Local URL is empty"
                ChatOpenAI(api_key=cfg.get("openai_key", "") or os.getenv("API_KEY"), base_url=url, model=model)
                return True, f"Connected to {url}"
            else:
                key = cfg.get("openai_key", "").strip()
                if not key:
                    return False, "OpenAI API key is empty"
                ChatOpenAI(api_key=key, model=model)
                return True, "API Connected"
        except Exception as e:
            return False, f"Connection failed: {e}"

    def on_config_updated(self, success, cfg):
        self.loaded_config = cfg
        if success:
            theme = cfg.get("theme")
            if theme:
                qss_path = os.path.join(self.project_root, "QSS", theme)
                if os.path.exists(qss_path):
                    try:
                        with open(qss_path, "r", encoding="utf-8") as fh:
                            QApplication.instance().setStyleSheet(fh.read())
                    except Exception:
                        pass
                self.current_theme = theme
            if cfg.get("api_mode") == "local":
                url = cfg.get("host_url", "")
                self.statusBar.showMessage(f"Connected to {url}")
            else:
                self.statusBar.showMessage("API Connected")
            self.chatSidebar.show_chat_ui()
        else:
            self.chatSidebar.show_not_configured()

    def _instantiate_llm(self):
        cfg = self.loaded_config or {}
        if not cfg:
            try:
                if os.path.exists(self.config_path):
                    with open(self.config_path, "r", encoding="utf-8") as f:
                        cfg = json.load(f)
                        self.loaded_config = cfg
                else:
                    QMessageBox.warning(self, "LLM", "No config.json found; please configure API in Settings.")
                    return None
            except Exception as e:
                QMessageBox.warning(self, "LLM", f"Failed to load config: {e}")
                return None

        if ChatOpenAI is None:
            QMessageBox.warning(self, "LLM", "ChatOpenAI import not available.")
            return None

        mode = cfg.get("api_mode", "online")
        model = cfg.get("model") or "gpt-4o"
        try:
            if mode == "local":
                url = cfg.get("host_url", "").strip()
                if not url:
                    QMessageBox.warning(self, "LLM", "Local URL empty in config.")
                    return None
                llm = ChatOpenAI(api_key=cfg.get("openai_key", "") or os.getenv("API_KEY"), base_url=url, model=model)
                return llm
            else:
                key = cfg.get("openai_key", "") or os.getenv("OPENAI_API_KEY") or os.getenv("LLM_KEY")
                if not key:
                    QMessageBox.warning(self, "LLM", "OpenAI API key empty in config or env.")
                    return None
                llm = ChatOpenAI(api_key=key, model=model)
                return llm
        except Exception as e:
            QMessageBox.warning(self, "LLM", f"Failed to instantiate LLM: {e}")
            return None

    def _call_llm(self, llm, human_content):
        try:
            from langchain.schema import HumanMessage
            msg = HumanMessage(content=human_content)
            res = llm.invoke([msg])
            if hasattr(res, "content"):
                return res.content
            if isinstance(res, dict) and "content" in res:
                return res["content"]
            gen = getattr(res, "generations", None)
            if gen:
                try:
                    return gen[0][0].text
                except Exception:
                    return str(gen)
            return str(res)
        except Exception as e:
            return f"LLM call failed: {e}"

    def generate_report(self):
        ef = self.evidence_file_path()
        if not os.path.exists(ef):
            QMessageBox.warning(self, "Report", "No evidence.json found. Connect a device to collect evidence first.")
            return
        try:
            with open(ef, "r", encoding="utf-8") as f:
                evidence_json = json.load(f)
        except Exception as e:
            QMessageBox.warning(self, "Report", f"Failed to load evidence.json: {e}")
            return

        llm = self._instantiate_llm()
        if llm is None:
            return

        self.statusBar.showMessage("Generating report...")
        sections = {}
        try:
            header = (
                "You are a seasoned mobile forensic analyst. "
                "Write a professional, well-structured, and detailed section (aim for 300-600 words) using:\n"
                "- concise paragraphs,\n- bullet lists where helpful,\n- clear assumptions/limitations, and\n- actionable insights.\n\n"
            )

            s1 = header + f"Section: Case Metadata\nEvidence:\n{json.dumps(evidence_json.get('device_info', {}), indent=2)}"
            sections["Case Metadata"] = self._call_llm(llm, s1)

            s2 = header + (
                "Section: File System Analysis\n"
                "Summarize notable file types, likely user behavior, anomalies, and suggested follow-ups.\n"
                f"Evidence:\n{json.dumps(evidence_json.get('files', []), indent=2)}"
            )
            sections["File System"] = self._call_llm(llm, s2)

            comm_ev = {"calls": evidence_json.get("calls", []), "sms": evidence_json.get("sms", [])}
            s3 = header + (
                "Section: Communications Analysis\n"
                "Discuss calling/SMS patterns, frequencies, timing clusters, and potential points of interest.\n"
                f"Evidence:\n{json.dumps(comm_ev, indent=2)}"
            )
            sections["Communications"] = self._call_llm(llm, s3)

            s4 = header + (
                "Section: Usage Stats Analysis\n"
                "Explain foreground/background app events, session bursts, and correlate with other artifacts.\n"
                f"Evidence:\n{json.dumps(evidence_json.get('usage_stats', []), indent=2)}"
            )
            sections["Usage Stats"] = self._call_llm(llm, s4)

            os.makedirs(os.path.join(self.project_root, "Docfiles"), exist_ok=True)
            case_id = evidence_json.get("device_info", {}).get("case_id", f"CASE-{datetime.now().strftime('%Y%m%d%H%M%S')}")
            out_fn = os.path.join(self.project_root, "Docfiles", f"report_{case_id}.docx")
            doc = Document()
            doc.add_heading(f"Forensic Report - {case_id}", 0)
            doc.add_heading("Case Metadata", level=1)
            doc.add_paragraph(sections.get("Case Metadata", ""))
            doc.add_heading("File System Analysis", level=1)
            doc.add_paragraph(sections.get("File System", ""))
            doc.add_heading("Communications Analysis", level=1)
            doc.add_paragraph(sections.get("Communications", ""))
            doc.add_heading("Usage Stats Analysis", level=1)
            doc.add_paragraph(sections.get("Usage Stats", ""))
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph("This is an automatically generated preliminary forensic report with expanded detail. Further manual review is recommended.")
            doc.save(out_fn)
            QMessageBox.information(self, "Report", f"Report saved to {out_fn}")
            self.statusBar.showMessage(f"Report generated: {out_fn}")
        except Exception as e:
            QMessageBox.warning(self, "Report", f"Failed to generate report: {e}")
            self.statusBar.showMessage("Report generation failed")


# ============================================================
# Entry point
# ============================================================

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = DroidForen()
    window.show()
    sys.exit(app.exec_())
